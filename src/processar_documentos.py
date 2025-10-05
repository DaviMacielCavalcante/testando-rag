#!/usr/bin/env python3
"""
processar_documentos.py
Script para processar documentos DOCX e adicionar ao ChromaDB
"""

import os
import chromadb
from sentence_transformers import SentenceTransformer
from pathlib import Path

def extrair_texto_docx(caminho_arquivo):
    """
    Extrai texto de arquivo DOCX
    """
    try:
        from docx import Document
        
        print(f"📄 Extraindo texto de: {Path(caminho_arquivo).name}")
        
        doc = Document(caminho_arquivo)
        texto = ""
        
        # Extrair parágrafos
        for paragrafo in doc.paragraphs:
            if paragrafo.text.strip():  # Só adicionar se não for vazio
                texto += paragrafo.text + "\n"
        
        # Extrair tabelas caso existam
        if doc.tables:
            linhas_tabela = []

            for tabela in doc.tables:
                for linha in tabela.rows:
                    linha_texto = " | ".join(celula.text for celula in linha.cells)
                    
                    if linha_texto.strip():
                        linhas_tabela.append(linha_texto + "\n")
        
        return texto.strip()
        
    except ImportError:
        print("python-docx não instalado! Execute: pip install python-docx")
        return None
        
    except Exception as e:
        print(f"Erro ao extrair texto de {caminho_arquivo}: {e}")
        return None


def dividir_texto_em_chunks(texto, tamanho=1000, sobreposicao=200):
    """
    Divide texto em chunks menores com sobreposição
    """
    if not texto or len(texto) < 50:
        return []
    
    chunks = []
    inicio = 0
    
    while inicio < len(texto):
        fim = inicio + tamanho
        
        # Se não é o último chunk, tentar quebrar em local natural
        if fim < len(texto):
    
            quebra_paragrafo = texto.rfind('\n\n', inicio, fim)
            if quebra_paragrafo > inicio + tamanho * 0.5:
                fim = quebra_paragrafo + 2
            
            quebra_frase = texto.rfind('. ', inicio, fim)
            if quebra_frase > inicio + tamanho * 0.7:
                fim = quebra_frase + 2
        
        chunk = texto[inicio:fim].strip()
        
        if len(chunk) > 50:  
            chunks.append(chunk)
            
        inicio = fim - sobreposicao if fim - sobreposicao >= 0 else 0
    
    return chunks


def processar_documentos():
    """
    Função principal para processar todos os documentos DOCX
    """
    print("🚀 PROCESSAMENTO DE DOCUMENTOS DOCX PARA CHROMADB")
    print("=" * 60)
    
    # ===== CONFIGURAÇÕES =====
    PASTA_DOCUMENTOS = "./documentos"  # 👈 Pasta com seus DOCX
    BASE_DADOS = "./chromadb_base"          # Onde salvar ChromaDB
    COLECAO = "documentos_word"             # Nome da coleção
    
    # Configurações de chunk
    CHUNK_SIZE = 1000      # Tamanho de cada pedaço
    CHUNK_OVERLAP = 200    # Sobreposição entre pedaços
    
    # ===== 1. VERIFICAR PASTA DE DOCUMENTOS =====
    
    if not os.path.exists(PASTA_DOCUMENTOS):
        print(f"❌ Pasta não encontrada: {PASTA_DOCUMENTOS}")
        print(f"💡 Crie a pasta e coloque seus arquivos .docx nela:")
        print(f"   mkdir {PASTA_DOCUMENTOS}")
        return False
    
    # Encontrar arquivos DOCX
    pasta = Path(PASTA_DOCUMENTOS)
    arquivos_docx = []
    
    # Buscar arquivos .docx e .DOCX
    for padrao in ["*.docx", "*.DOCX"]:
        arquivos_docx.extend(list(pasta.glob(padrao)))
    
    # Filtrar arquivos temporários do Word (começam com ~$)
    arquivos_docx = [f for f in arquivos_docx if not f.name.startswith("~$")]
    
    if not arquivos_docx:
        print(f"❌ Nenhum arquivo .docx encontrado em {PASTA_DOCUMENTOS}")
        print(f"💡 Coloque arquivos .docx na pasta e tente novamente")
        return False
    
    print(f"📁 Encontrados {len(arquivos_docx)} arquivos DOCX:")
    for arquivo in arquivos_docx:
        tamanho_mb = arquivo.stat().st_size / (1024 * 1024)
        print(f"  📄 {arquivo.name} ({tamanho_mb:.1f} MB)")
    
    # ===== 2. CONFIGURAR CHROMADB =====
    
    print(f"\n⚙️ Configurando ChromaDB em {BASE_DADOS}...")
    
    try:
        client = chromadb.PersistentClient(path=BASE_DADOS)
        
        # Criar ou conectar à coleção
        try:
            collection = client.get_collection(COLECAO)
            print(f"✓ Conectado à coleção existente '{COLECAO}'")
            print(f"  📊 Documentos atuais: {collection.count()}")
        except:
            collection = client.create_collection(COLECAO)
            print(f"✓ Nova coleção '{COLECAO}' criada")
        
    except Exception as e:
        print(f"❌ Erro ao configurar ChromaDB: {e}")
        return False
    
    # ===== 3. CARREGAR MODELO DE EMBEDDINGS =====
    
    print(f"\n🔢 Carregando modelo de embeddings...")
    
    try:
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✓ Modelo de embeddings carregado")
    except Exception as e:
        print(f"❌ Erro ao carregar modelo: {e}")
        print("💡 Execute: pip install sentence-transformers")
        return False
    
    # ===== 4. PROCESSAR CADA ARQUIVO =====
    
    print(f"\n📚 Processando documentos...")
    print("=" * 40)
    
    total_chunks_adicionados = 0
    arquivos_processados = 0
    
    for i, arquivo_path in enumerate(arquivos_docx, 1):
        print(f"\n[{i}/{len(arquivos_docx)}] Processando: {arquivo_path.name}")
        
        try:
            # Extrair texto
            texto = extrair_texto_docx(str(arquivo_path))
            
            if not texto:
                print(f"⚠️ Nenhum texto extraído de {arquivo_path.name}")
                continue
            
            print(f"✓ Texto extraído: {len(texto):,} caracteres")
            
            # Dividir em chunks
            chunks = dividir_texto_em_chunks(texto, CHUNK_SIZE, CHUNK_OVERLAP)
            
            if not chunks:
                print(f"⚠️ Nenhum chunk criado para {arquivo_path.name}")
                continue
                
            print(f"✓ Dividido em {len(chunks)} chunks")
            
            # Criar metadados para cada chunk
            metadados = []
            for j, chunk in enumerate(chunks):
                metadados.append({
                    "arquivo": arquivo_path.name,
                    "caminho_completo": str(arquivo_path.absolute()),
                    "chunk_index": j,
                    "total_chunks": len(chunks),
                    "tamanho_chunk": len(chunk),
                    "palavras_chunk": len(chunk.split()),
                    "tamanho_arquivo_original": len(texto)
                })
            
            # Gerar IDs únicos
            base_id = arquivo_path.stem  # Nome sem extensão
            ids = [f"{base_id}_chunk_{j:03d}" for j in range(len(chunks))]
            
            # Verificar se já existem chunks deste arquivo
            try:
                existentes = collection.get(ids=ids)
                if existentes['ids']:
                    print(f"⚠️ Arquivo já processado anteriormente. Pulando...")
                    continue
            except:
                pass  # Continuar se não conseguir verificar
            
            # Criar embeddings
            print("🔄 Criando embeddings...")
            embeddings = embedding_model.encode(chunks, show_progress_bar=False).tolist()
            
            # Adicionar ao ChromaDB
            collection.add(
                documents=chunks,
                embeddings=embeddings,
                metadatas=metadados,
                ids=ids
            )
            
            print(f"✅ {len(chunks)} chunks adicionados ao ChromaDB")
            total_chunks_adicionados += len(chunks)
            arquivos_processados += 1
            
        except Exception as e:
            print(f"❌ Erro ao processar {arquivo_path.name}: {e}")
            continue
    
    # ===== 5. RESUMO FINAL =====
    
    print(f"\n" + "=" * 60)
    print(f"🎉 PROCESSAMENTO CONCLUÍDO!")
    print(f"=" * 60)
    print(f"📊 Arquivos processados: {arquivos_processados}/{len(arquivos_docx)}")
    print(f"📦 Total de chunks criados: {total_chunks_adicionados}")
    print(f"💾 Base de dados salva em: {BASE_DADOS}")
    print(f"📚 Coleção: {COLECAO}")
    
    # Estatísticas finais
    total_final = collection.count()
    print(f"📈 Total de chunks na base: {total_final}")
    
    if total_chunks_adicionados > 0:
        print(f"\n✅ Sistema pronto para uso!")
        print(f"💡 Próximos passos:")
        print(f"   1. Configure um modelo LLM (Ollama, OpenAI, etc)")
        print(f"   2. Execute o sistema RAG")
        print(f"   3. Faça perguntas sobre seus documentos!")
        
        # Teste rápido de busca
        print(f"\n🔍 Teste rápido de busca...")
        try:
            query_test = "principais"
            query_embedding = embedding_model.encode(query_test).tolist()
            resultados = collection.query(
                query_embeddings=[query_embedding],
                n_results=2
            )
            
            if resultados['documents'][0]:
                print(f"✓ Busca funcionando! Encontrados documentos relevantes")
                print(f"  Exemplo: {resultados['documents'][0][0][:100]}...")
            else:
                print("⚠️ Busca não retornou resultados")
                
        except Exception as e:
            print(f"⚠️ Erro no teste de busca: {e}")
    
    else:
        print(f"\n⚠️ Nenhum documento foi processado com sucesso")
        print(f"💡 Verifique:")
        print(f"   - Os arquivos são .docx válidos?")
        print(f"   - Os arquivos não estão corrompidos?")
        print(f"   - Você tem permissão para ler os arquivos?")
    
    return total_chunks_adicionados > 0


def verificar_status():
    """
    Verifica o status atual da base de dados
    """
    BASE_DADOS = "./chromadb_base"
    COLECAO = "documentos_word"
    
    print("📊 STATUS DA BASE DE DADOS")
    print("=" * 40)
    
    try:
        if not os.path.exists(BASE_DADOS):
            print("❌ Base de dados não existe ainda")
            print("💡 Execute o processamento primeiro")
            return
        
        client = chromadb.PersistentClient(path=BASE_DADOS)
        collection = client.get_collection(COLECAO)
        
        total = collection.count()
        print(f"✅ Base de dados encontrada")
        print(f"📦 Total de chunks: {total}")
        
        if total > 0:
            # Pegar alguns exemplos
            amostra = collection.get(limit=3)
            
            # Contar arquivos únicos
            arquivos = set()
            for metadata in amostra['metadatas']:
                arquivos.add(metadata.get('arquivo', 'Desconhecido'))
            
            print(f"📚 Arquivos processados: {len(arquivos)}")
            print(f"📄 Exemplos:")
            for arquivo in list(arquivos):
                print(f"   - {arquivo}")
                
            print(f"\n✅ Sistema pronto para uso!")
        else:
            print(f"⚠️ Base existe mas está vazia")
            
    except Exception as e:
        print(f"❌ Erro ao verificar base: {e}")


def menu_principal():
    """
    Menu principal do script
    """
    while True:
        print("\n" + "=" * 50)
        print("📚 PROCESSAMENTO DE DOCUMENTOS DOCX")
        print("=" * 50)
        print("1. 📄 Processar documentos DOCX")
        print("2. 📊 Ver status da base de dados")
        print("3. 🗑️ Limpar base de dados")
        print("4. ❓ Ajuda")
        print("5. 🚪 Sair")
        print("-" * 50)
        
        escolha = input("Escolha uma opção (1-5): ").strip()
        
        if escolha == "1":
            sucesso = processar_documentos()
            if sucesso:
                print("\n✅ Processamento concluído com sucesso!")
            else:
                print("\n❌ Processamento falhou. Verifique os erros acima.")
                
        elif escolha == "2":
            verificar_status()
            
        elif escolha == "3":
            import shutil
            confirma = input("\n⚠️ Tem certeza que quer apagar a base? (sim/não): ")
            if confirma.lower() in ['sim', 's', 'yes', 'y']:
                try:
                    if os.path.exists("./chromadb_base"):
                        shutil.rmtree("./chromadb_base")
                        print("✅ Base de dados apagada!")
                    else:
                        print("ℹ️ Base de dados não existe")
                except Exception as e:
                    print(f"❌ Erro ao apagar: {e}")
            else:
                print("Operação cancelada")
                
        elif escolha == "4":
            print("\n📖 AJUDA")
            print("-" * 30)
            print("1. Coloque arquivos .docx na pasta 'meus_documentos'")
            print("2. Execute a opção 1 para processar")
            print("3. Use a opção 2 para verificar o status")
            print("4. Depois configure o sistema RAG para fazer perguntas")
            print("\nDependências necessárias:")
            print("pip install chromadb sentence-transformers python-docx")
            
        elif escolha == "5":
            print("\n👋 Até logo!")
            break
            
        else:
            print("\n❌ Opção inválida. Tente novamente.")


if __name__ == "__main__":
    # Verificar dependências básicas
    try:
        import chromadb
        import sentence_transformers
        import docx
    except ImportError as e:
        print("❌ Dependência não encontrada!")
        print("Execute: pip install chromadb sentence-transformers python-docx")
        exit(1)
    
    # Executar menu principal
    menu_principal()